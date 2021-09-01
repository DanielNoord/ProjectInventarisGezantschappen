#!/usr/bin/env python3

import re

import docx


def extract_persons_and_identifiers(filename):
    """Loads people mentioned and given identifier from specified file

    Args:
        filename (str): File to be scanned

    Returns:
        [[], [], [], ...] List of lists
            [string, string] Each list contains two strings which are:
                string: Data in inventory
                string: Identifier given to that data
    """
    doc = docx.Document(filename)
    all_text = []
    for para in doc.paragraphs:
        if para.text == "":
            break
        all_text.append(para.text.split("\n"))
    return all_text


def extract_persons(filename):
    """Loads people mentioned from "database" file

    Args:
        filename (str): File to be scanned

    Returns:
        list[str, str, ...]: List of persons (in string format) found in the file
    """
    regex = re.compile(r".*Personen:\n(-.*;)\n*Instituten:", re.DOTALL)
    doc = docx.Document(filename)
    all_text = []
    impossible = -1
    missing_possible = -1
    for para in doc.paragraphs:
        if len(para.runs) > 1:
            if para.runs[1].font.highlight_color == 6 or re.match(r"- .*?; 2", para.text):
                impossible += 1
            if para.runs[1].font.highlight_color == 7 or re.match(r"- .*?; 3", para.text):
                missing_possible += 1
        all_text.append(para.text)
    full_text = "\n".join(all_text)
    person_section = re.findall(regex, full_text)
    persons_in_file = person_section[0].split("\n")

    print(
        f"\nFound {len(persons_in_file)} people\n\
    Found {impossible} impossible entries and {missing_possible} missing but possible entries\n\
    About {((impossible + missing_possible) / len(persons_in_file)) * 100 // 1}% of entries\n"
    )

    return persons_in_file


def extract_translations(filename):
    """Loads translations from file

    Args:
        filename (str): File to be scanned

    Returns:
        dict{{}, {}, {}, ... }: Dictionary of translations (which are also dictionaries)
    """
    doc = docx.Document(filename)
    all_translations = {}
    for para in doc.paragraphs:
        translations = para.text.split("\n")
        translation = {
            "en_GB": translations[2],
            "it_IT": translations[1],
            "nl_NL": translations[0],
        }
        for line in translations[3:]:
            if line.startswith("Opmerking: "):
                translation["comment"] = translations[3][11:]  # Cuts of "Opmerking: "
            if line.startswith("position: "):
                translation["position"] = translations[3][10:]  # Cuts of "position: "

        all_translations[translations[0]] = translation
    return all_translations
