import json

import docx  # type: ignore[import-not-found]


def search_koelman(filename: str) -> None:
    """Checks all names in a database for hits in biographic register of Poelman book of Asker.

    Args:
        filename: al database
    """
    with open(filename, encoding="utf-8") as file:
        persons = json.load(file)
    del persons["$schema"]
    doc = docx.Document("koelman/Biografisch register NL Koelman.docx")
    names_register = [i.text.split(" ")[0].replace(",", "") for i in doc.paragraphs]

    for data in persons.values():
        if data["surname"] in names_register:
            print("\n", data["name"], data["surname"])
            print(doc.paragraphs[names_register.index(data["surname"])].text)

    print("Finished checking for hits in Koelman database!")


if __name__ == "__main__":
    search_koelman("inputs/Individuals.json")
