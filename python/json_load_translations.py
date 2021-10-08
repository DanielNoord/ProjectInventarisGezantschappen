#!/usr/bin/env python3

import json
import re

import docx  # type: ignore

from typing_utils import TranslationDict, TranslationDictTitles
from write_files import write_single_json_file


def load_translations_documents(filename: str) -> None:
    """Load translations from .docx and write .json for documents file

    Args:
        filename (str): Filename of the input file
    """
    doc = docx.Document(filename)
    all_documents: TranslationDict = {
        "$schema": "../../static/JSON/DocumentTitles.json"
    }

    for para in doc.paragraphs:
        (it_it_trans, nl_nl_trans, en_gb_trans) = re.split(r"\n.*?", para.text)

        all_documents[it_it_trans] = {"en_GB": en_gb_trans, "nl_NL": nl_nl_trans}

    write_single_json_file(all_documents, "outputs", "DocumentTitles.json")


def load_translations_functions(filename: str) -> None:
    """Load translations from .docx and write .json for functions file

    Args:
        filename (str): Filename of the input file
    """
    doc = docx.Document(filename)
    all_functions: TranslationDict = {"$schema": "../../static/JSON/Functions.json"}

    for para in doc.paragraphs:
        (nl_nl_trans, it_it_trans, en_gb_trans) = re.split(r"\n.*?", para.text)

        all_functions[it_it_trans] = {
            "en_GB": en_gb_trans,
            "nl_NL": nl_nl_trans,
        }

    write_single_json_file(all_functions, "outputs", "Functions.json")


def load_translations_placenames(filename: str) -> None:
    """Load translations from .docx and write .json for placename file

    Args:
        filename (str): Filename of the input file
    """
    doc = docx.Document(filename)
    all_placenames: TranslationDict = {"$schema": "../../static/JSON/Placenames.json"}

    for para in doc.paragraphs:
        (it_it_trans, nl_nl_trans, en_gb_trans) = re.split(r"\n.*?", para.text)

        all_placenames[it_it_trans] = {"en_GB": en_gb_trans, "nl_NL": nl_nl_trans}

    write_single_json_file(all_placenames, "outputs", "Placenames.json")


def load_translations_titles(filename: str) -> None:
    """Load translations from .docx and write .json for titles file

    Args:
        filename (str): Filename of the input file
    """
    doc = docx.Document(filename)
    all_titles: TranslationDictTitles = {"$schema": "../../static/JSON/Titles.json"}

    for para in doc.paragraphs:
        (nl_nl_trans, it_it_trans, en_gb_trans, position) = re.split(
            r"\n.*?", para.text
        )

        all_titles[it_it_trans] = {
            "en_GB": en_gb_trans,
            "nl_NL": nl_nl_trans,
            "position": position.replace("position: ", ""),
        }

    write_single_json_file(all_titles, "outputs", "Titles.json")


if __name__ == "__main__":
    load_translations_documents("inputs/Translations/TranslatedDocumentTitles.docx")
    # load_translations_functions("inputs/Translations/TranslatedFunctions.docx")
    # load_translations_placenames("inputs/Translations/TranslatedPlacenames.docx")
    # load_translations_titles("inputs/Translations/TranslatedTitles.docx")
