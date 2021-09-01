#!/usr/bin/env python3

import os
import re

import docx
from docx.shared import Pt


def list_to_be_translated(list_to_write: list[str], output_name: str) -> None:
    """Write a docx file based on input. Adds two white lines after every entry

    Args:
        list_to_write (list[str]): List of translated names to be written, all in strings
        output_name (str): Name of the output file
    """
    out_doc = docx.Document()
    for line in list_to_write:
        paragraph = out_doc.add_paragraph(f"{line}\n\n")

        # Add indentation
        paragraph_format = paragraph.paragraph_format
        paragraph_format.first_line_indent = Pt(-10)

    # Check if outputs directory exists and then write file
    os.makedirs(os.path.join(os.getcwd(), r"outputs"), exist_ok=True)
    out_doc.save(f"outputs/Translated{output_name}.docx")
    print(f"Wrote file at outputs/Translated{output_name}.docx")


def list_of_translated_data(list_to_write: list[list], output_name: str) -> None:
    """Write a docx file based on input

    Args:
        list_to_write (list[list]): List of translated data to be written, all in strings.
            Follows pattern [Data in inventaris, identifier, data in IT, data in NL, data in EN]
        output_name (str): Name of the output file
    """
    out_doc = docx.Document()
    for entry in list_to_write:
        paragraph = out_doc.add_paragraph("")
        run = paragraph.add_run(entry[0])
        run.bold = True
        for line in entry[1:]:
            paragraph.add_run("\n")
            run = paragraph.add_run(line)

        # Add indentation
        paragraph_format = paragraph.paragraph_format
        paragraph_format.first_line_indent = Pt(-10)

    # Check if outputs directory exists and then write file
    os.makedirs(os.path.join(os.getcwd(), r"outputs"), exist_ok=True)
    out_doc.save(f"outputs/Translated{output_name}.docx")
    print(f"Wrote file at outputs/Translated{output_name}.docx")


def list_of_translated_data_with_style(list_to_write: list[list], output_name: str) -> None:
    """Write a docx file based on input with style applied

    Args:
        list_to_write (list[list]): List of translated data to be written, all in strings.
            Follows pattern [Data in inventaris, identifier, data in IT, data in NL, data in EN]
        output_name (str): Name of the output file
    """
    out_doc = docx.Document()
    for entry in list_to_write:
        paragraph = out_doc.add_paragraph("")
        run = paragraph.add_run(entry[0])
        run.bold = True
        for line in entry[1:]:
            # Check for translation (denotade by {translation})
            line = line.replace("{", "(").replace("}", ")")
            # Check for parts that should be italic (denoted by _italian_)
            split_lines = re.split(r"(_.*?_)", line)

            paragraph.add_run("\n")
            for group in split_lines:
                if group != "" and group[0] == "_":
                    run = paragraph.add_run(group[1:-1])
                    run.italic = True
                else:
                    run = paragraph.add_run(group)

        # Add indentation
        paragraph_format = paragraph.paragraph_format
        paragraph_format.first_line_indent = Pt(-10)

    # Check if outputs directory exists and then write file
    os.makedirs(os.path.join(os.getcwd(), r"outputs"), exist_ok=True)
    out_doc.save(f"outputs/TranslatedStylized{output_name}.docx")
    print(f"Wrote file at outputs/TranslatedStylized{output_name}.docx")


def list_with_style(list_of_lines: list[str], output_name: str) -> None:
    """Write a docx file based on input

    Args:
        list_of_lines (list[str]): List of lines to be written
        output_name (str): Name of output file
    """
    out_doc = docx.Document()
    for line in list_of_lines:
        # Check for translation (denotade by {translation})
        line = line.replace("{", "(").replace("}", ")")
        # Check for parts that should be italic (denoted by _italian_)
        split_lines = re.split(r"(_.*?_)", line)

        paragraph = out_doc.add_paragraph("")
        for group in split_lines:
            if group != "" and group[0] == "_":
                run = paragraph.add_run(group[1:-1])
                run.italic = True
            else:
                run = paragraph.add_run(group)
        # Add indentation
        paragraph_format = paragraph.paragraph_format
        paragraph_format.first_line_indent = Pt(-10)

    # Check if outputs directory exists and then write file
    os.makedirs(os.path.join(os.getcwd(), r"outputs"), exist_ok=True)
    out_doc.save(f"outputs/Stylized{output_name}.docx")
    print(f"Wrote file at outputs/Stylized{output_name}.docx")


def database(dict_of_individuals: dict, output_name: str, skip_types: list[int]) -> None:
    """Write a docx file based on input dictionary

    Args:
        dict_of_individuals (dict): Dictionary containing all individuals
        output_name (str): Name of output file
        skip_types (list[int]): Number of person_types to skip (for example [1, 2])
    """
    if skip_types:
        output_name = f"{output_name}_without_types_{','.join(str(i) for i in skip_types)}"

    amount = 0
    out_doc = docx.Document()

    for identifier, data in dict_of_individuals.items():
        if data["person_type"] not in skip_types:
            paragraph = out_doc.add_paragraph("")
            run = paragraph.add_run(f"{identifier}\n")
            run.bold = True
            paragraph.add_run(f"Type: {str(data['person_type'])}\n")
            paragraph.add_run(f"Surname: {data['surname']}\n")
            paragraph.add_run(f"Name: {data['name']}\n")
            paragraph.add_run(f"Date of birth: {data['date_of_birth']}\n")
            paragraph.add_run(f"Place of birth: {data['place_of_birth']}\n")
            paragraph.add_run(f"Date of death: {data['date_of_death']}\n")
            paragraph.add_run(f"Place of death: {data['place_of_death']}\n")
            title_string = f"Titles: {'| '.join(f'{i} ({j})' for i, j in data['titles'])}\n"
            paragraph.add_run(title_string.replace(" (None)", ""))
            func_string = f"Functions: {'| '.join(f'{i} ({j})' for i, j in data['functions'])}\n"
            paragraph.add_run(func_string.replace(" (None)", ""))
            paragraph.add_run(f"Comment: {data['comment']}\n")
            paragraph.add_run(f"Comment from Daniël: {data['comment_daniel']}\n")
            paragraph.add_run(f"Sources: {'| '.join(data['sources'])}\n")
            paragraph.add_run(f"Images: {'| '.join(data['images'])}\n")

            paragraph_format = paragraph.paragraph_format
            paragraph_format.first_line_indent = Pt(-10)
            amount += 1

    # Check if outputs directory exists and then write file
    os.makedirs(os.path.join(os.getcwd(), r"outputs"), exist_ok=True)
    out_doc.save(f"outputs/{output_name}.docx")
    print(f"Wrote file at outputs/{output_name}.docx with {amount} people")
