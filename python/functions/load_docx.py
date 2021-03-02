import re

import docx


def extract_persons_and_identifiers(filename):
    """Loads people mentioned and given identifier from specified file

    Args:
        filename (str): File to be scanned

    Returns:
        [[], [], [], ...] List of lists
            [string, string] Each list contains two strings which are:
                string: Data in inventais
                string: Identifier given to that data
    """
    doc = docx.Document(filename)
    all_text = []
    for para in doc.paragraphs:
        if para.text == "":
            break
        all_text.append(para.text.split("\n"))
    return all_text


def extract_persons(filename):
    """Loads people mentioned from "database" file

    Args:
        filename (str): File to be scanned

    Returns:
        list[str, str, ...]: List of persons (in string format) found in the file
    """
    regex = re.compile(r".*Personen:\n(-.*;)\n*Instituten:", re.DOTALL)
    doc = docx.Document(filename)
    all_text = []
    count_red = -1
    count_yellow = -1
    for para in doc.paragraphs:
        if len(para.runs) > 1:
            if para.runs[1].font.highlight_color == 6:
                count_red += 1
            if para.runs[1].font.highlight_color == 7:
                count_yellow += 1
        all_text.append(para.text)
    full_text = "\n".join(all_text)
    person_section = re.findall(regex, full_text)
    persons_in_file = person_section[0].split("\n")

    print(
        f"Found {len(persons_in_file)} people\n\
    Found {count_red} red entries and {count_yellow} yellow entries\n\
    ({((count_red + count_yellow) / len(persons_in_file)) * 100 // 1}%)"
    )

    return persons_in_file


def extract_volumes(filename):
    """Loads volumes from "fondo" file

    Args:
        filename (str): File to be scanned

    Returns:
        list[str, str, ...]: List of volumes (in string format) found in the file
    """
    regex = re.compile(r".*?(Volume\n.*?)(?=Volume|$)", re.DOTALL)
    doc = docx.Document(filename)
    all_text = []
    for para in doc.paragraphs:
        all_text.append(para.text)
    full_text = "\n".join(all_text)
    return re.findall(regex, full_text)
