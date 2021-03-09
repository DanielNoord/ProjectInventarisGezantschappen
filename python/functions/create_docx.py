import os
import re

import docx
from docx.shared import Pt


def list_to_be_translated(list_to_write, output_name):
    """Write a docx file based on input. Adds two white lines after every entry.

    Args:
        list_to_write (list): List of translated names to be written, all in strings.
    """
    out_doc = docx.Document()
    for line in list_to_write:
        paragraph = out_doc.add_paragraph(f"{line}\n\n")

        # Add indentation
        paragraph_format = paragraph.paragraph_format
        paragraph_format.first_line_indent = Pt(-10)

    # Check if outputs directory exists and then write file
    os.makedirs(os.path.join(os.getcwd(), r"outputs"), exist_ok=True)
    out_doc.save(f"outputs/Translated{output_name}.docx")
    print(f"Wrote file at outputs/Translated{output_name}.docx")


def list_of_translated_data(list_to_write, output_name):
    """Write a docx file based on input

    Args:
        list_to_write (list): List of translated data to be written, all in strings.
                Follows pattern [Data in inventaris, identifier, data in IT, data in NL, data in EN]
    """
    out_doc = docx.Document()
    for entry in list_to_write:
        paragraph = out_doc.add_paragraph("")
        run = paragraph.add_run(entry[0])
        run.bold = True
        for line in entry[1:]:
            paragraph.add_run("\n")
            run = paragraph.add_run(line)

        # Add indentation
        paragraph_format = paragraph.paragraph_format
        paragraph_format.first_line_indent = Pt(-10)

    # Check if outputs directory exists and then write file
    os.makedirs(os.path.join(os.getcwd(), r"outputs"), exist_ok=True)
    out_doc.save(f"outputs/Translated{output_name}.docx")
    print(f"Wrote file at outputs/Translated{output_name}.docx")

def list_of_translated_data_with_style(list_to_write, output_name):
    """Write a docx file based on input with style applied

    Args:
        list_to_write (list): List of translated data to be written, all in strings.
                Follows pattern [Data in inventaris, identifier, data in IT, data in NL, data in EN]
    """
    out_doc = docx.Document()
    for entry in list_to_write:
        paragraph = out_doc.add_paragraph("")
        run = paragraph.add_run(entry[0])
        run.bold = True
        for line in entry[1:]:
            # Check for translation (denotade by {translation})
            line = line.replace("{", "(").replace("}", ")")
            # Check for parts that should be italic (denoted by _italian_)
            split_lines = re.split(r"(_.*?_)", line)

            paragraph.add_run("\n")
            for group in split_lines:
                if group != "" and group[0] == "_":
                    run = paragraph.add_run(group[1:-1])
                    run.italic = True
                else:
                    run = paragraph.add_run(group)

        # Add indentation
        paragraph_format = paragraph.paragraph_format
        paragraph_format.first_line_indent = Pt(-10)

    # Check if outputs directory exists and then write file
    os.makedirs(os.path.join(os.getcwd(), r"outputs"), exist_ok=True)
    out_doc.save(f"outputs/TranslatedStylized{output_name}.docx")
    print(f"Wrote file at outputs/TranslatedStylized{output_name}.docx")


def list_with_style(list_of_lines, output_name):
    """Write a docx file based on input

    Args:
        list_of_lines (list): List of lines to be written
        filename (str): Name of output file
    """
    out_doc = docx.Document()
    for line in list_of_lines:
        # Check for translation (denotade by {translation})
        line = line.replace("{", "(").replace("}", ")")
        # Check for parts that should be italic (denoted by _italian_)
        split_lines = re.split(r"(_.*?_)", line)

        paragraph = out_doc.add_paragraph("")
        for group in split_lines:
            if group != "" and group[0] == "_":
                run = paragraph.add_run(group[1:-1])
                run.italic = True
            else:
                run = paragraph.add_run(group)
        # Add indentation
        paragraph_format = paragraph.paragraph_format
        paragraph_format.first_line_indent = Pt(-10)

    # Check if outputs directory exists and then write file
    os.makedirs(os.path.join(os.getcwd(), r"outputs"), exist_ok=True)
    out_doc.save(f"outputs/Stylized{output_name}.docx")
    print(f"Wrote file at outputs/Stylized{output_name}.docx")

def database(dict_of_individuals, output_name, skip_types):
    """Write a docx file based on input dictionary

    Args:
        dict_of_individuals (dict): Dictionary containing all individuals
        output_name (str): Name of output file
        skip_types (list): Number of person_types to skip (for example [1, 2])
    """
    if skip_types:
        output_name = f"{output_name}_without_types_{','.join(skip_types)}"
    out_doc = docx.Document()
    for identifier, data in dict_of_individuals.items():
        if data["person_type"] not in skip_types:
            paragraph = out_doc.add_paragraph("")
            run = paragraph.add_run(f"{identifier}\n")
            run.bold = True
            paragraph.add_run(f"Type: {str(data['person_type'])}\n")
            paragraph.add_run(f"Surname: {data['surname']}\n")
            paragraph.add_run(f"Name: {data['name']}\n")
            paragraph.add_run(f"Nationality: {data['nationality']}\n")
            paragraph.add_run(f"Titles: {'| '.join(data['titles'])}\n")
            func_string = f"Functions: {'| '.join(f'{i} ({j})' for i, j in data['functions'])}\n"
            paragraph.add_run(func_string.replace(" (None)", ""))
            paragraph.add_run(f"Place of residence: {data['place of residence']}\n")
            paragraph.add_run(f"Comment: {data['comment']}\n")
            paragraph.add_run(f"Sources: {'| '.join(data['sources'])}\n")

            paragraph_format = paragraph.paragraph_format
            paragraph_format.first_line_indent = Pt(-10)



    # Check if outputs directory exists and then write file
    os.makedirs(os.path.join(os.getcwd(), r"outputs"), exist_ok=True)
    out_doc.save(f"outputs/{output_name}.docx")
    print(f"Wrote file at outputs/{output_name}.docx")
