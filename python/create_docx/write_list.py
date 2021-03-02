import os
import re
import docx
from docx.shared import Pt

def write_list(list_of_lines, filename):
    """ Write a docx file based on input

    Args:
        list_of_lines (list): List of lines to be written
        filename (str): Name of output file
    """
    doc = docx.Document()
    for line in list_of_lines:
        # Check for translation (denotade by {translation})
        line = line.replace("{", "(").replace("}", ")")
        # Check for parts that should be italic (denoted by _italian_)
        split_lines = re.split(r"(_.*?_)", line)

        paragraph = doc.add_paragraph('')
        for group in split_lines:
            if group != "" and group[0] == "_":
                run = paragraph.add_run(group[1:-1])
                run.italic = True
            else:
                run = paragraph.add_run(group)
        # Add indentation
        paragraph_format = paragraph.paragraph_format
        paragraph_format.first_line_indent = Pt(-10)

    # Check if outputs directory exists and then write file
    os.makedirs(os.path.join(os.getcwd(), r'outputs'), exist_ok=True)
    doc.save(filename)