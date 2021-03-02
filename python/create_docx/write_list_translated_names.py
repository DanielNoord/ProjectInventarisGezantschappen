import os
import re
import docx
from docx.shared import Pt


def write_list_translated_names(list_to_write, output_name):
    """ Write a docx file based on input

    Args:
        list_to_write (list): List of translated names to be written, all in strings.
                Follows pattern [Data in inventaris, identifier, name in IT, name in NL, name in EN]
    """
    out_doc = docx.Document()
    for entry in list_to_write:
        paragraph = out_doc.add_paragraph("")
        run = paragraph.add_run(entry[0])
        run.bold = True
        for line in entry[1:]:
            paragraph.add_run("\n")
            # Check for translation (denotade by {translation})
            line = line.replace("{", "(").replace("}", ")")
            # Check for parts that should be italic (denoted by _italian_)
            split_line = re.split(r"(_.*?_)", line)

            for group in split_line:
                if group != "" and group[0] == "_":
                    run = paragraph.add_run(group[1:-1])
                    run.italic = True
                else:
                    run = paragraph.add_run(group)

        # Add indentation
        paragraph_format = paragraph.paragraph_format
        paragraph_format.first_line_indent = Pt(-10)

    # Check if outputs directory exists and then write file
    os.makedirs(os.path.join(os.getcwd(), r'outputs'), exist_ok=True)
    out_doc.save(f"outputs/Translated{output_name}.docx")
